import csv
import hashlib
import json
import math
from io import StringIO
from pathlib import Path
from typing import Any

from sqlalchemy import bindparam, text
from sqlalchemy.dialects.postgresql import JSONB

from app.services.secrets import resolve_openai_api_key

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv", ".pdf", ".docx"}
DEFAULT_CHUNK_SIZE = 1200
DEFAULT_CHUNK_OVERLAP = 160
DEFAULT_EMBEDDING_DIM = 1536
LOCAL_EMBEDDING_MODELS = {"local", "local-hash", "local-embedding", "local-keyword", "keyword"}


class DocumentProcessingError(RuntimeError):
    def __init__(self, stage: str, message: str) -> None:
        super().__init__(message)
        self.stage = stage
        self.message = message


def extract_text_from_file(path: str | Path, file_type: str | None = None) -> str:
    file_path = Path(path)
    if not file_path.exists():
        raise DocumentProcessingError("parse", f"file not found: {file_path}")

    suffix = file_path.suffix.lower()
    try:
        if suffix == ".md" or _is_markdown_content_type(file_type):
            return _extract_markdown_text(file_path)
        if suffix == ".txt" or _is_text_content_type(file_type):
            return _extract_plain_text(file_path)
        if suffix == ".json" or file_type == "application/json":
            return _extract_json_text(file_path)
        if suffix == ".csv" or file_type in {"text/csv", "application/csv"}:
            return _extract_csv_text(file_path)
        if suffix == ".pdf" or file_type == "application/pdf":
            return _extract_pdf_text(file_path)
        if suffix == ".docx" or _is_docx_content_type(file_type):
            return _extract_docx_text(file_path)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        message = f"failed to parse {suffix or file_type}: {exc}"
        raise DocumentProcessingError("parse", message) from exc

    raise DocumentProcessingError(
        "parse",
        f"unsupported document format: {suffix or file_type or 'unknown'}",
    )


def chunk_text(
    text_value: str,
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[str]:
    text_value = _normalize_text(text_value)
    if not text_value:
        return []
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap must be greater than or equal to 0 and less than chunk_size")

    chunks: list[str] = []
    start = 0
    while start < len(text_value):
        end = min(start + chunk_size, len(text_value))
        if end < len(text_value):
            boundary = max(
                text_value.rfind("\n\n", start, end),
                text_value.rfind("\n", start, end),
                text_value.rfind(" ", start, end),
            )
            if boundary > start + chunk_size // 2:
                end = boundary
        chunk = text_value[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text_value):
            break
        start = max(end - overlap, 0)
    return chunks


def rank_chunks(chunks: list[dict[str, Any]], query: str) -> list[dict[str, Any]]:
    query_lower = query.lower().strip()
    query_terms = {term for term in query_lower.split() if term}
    ranked: list[dict[str, Any]] = []
    for chunk in chunks:
        content = str(chunk.get("content") or "")
        content_lower = content.lower()
        if not query_terms:
            score = 0.0
        else:
            matched = sum(1 for term in query_terms if term in content_lower)
            score = matched / len(query_terms)
        if score == 0.0 and query_lower and query_lower in content_lower:
            score = 1.0
        ranked.append({**chunk, "score": score})
    return sorted(ranked, key=lambda item: (item["score"], item.get("id") or 0), reverse=True)


async def retrieve_chunks(
    conn,
    *,
    knowledge_base_id: int,
    query: str,
    top_k: int,
    score_threshold: float,
) -> list[dict[str, Any]]:
    knowledge_base = await _get_knowledge_base_for_processing(conn, knowledge_base_id)
    vector_results = await _retrieve_chunks_by_vector(
        conn,
        knowledge_base=knowledge_base,
        query=query,
        top_k=top_k,
        score_threshold=score_threshold,
    )
    if vector_results:
        return vector_results
    return await _retrieve_chunks_by_keyword(
        conn,
        knowledge_base_id=knowledge_base_id,
        query=query,
        top_k=top_k,
        score_threshold=score_threshold,
    )


async def _retrieve_chunks_by_vector(
    conn,
    *,
    knowledge_base: dict[str, Any],
    query: str,
    top_k: int,
    score_threshold: float,
) -> list[dict[str, Any]]:
    try:
        query_embedding = (await embed_texts(conn, knowledge_base, [query]))[0]
    except Exception:
        return []

    result = await conn.execute(
        text(
            """
            SELECT
              c.id,
              c.content,
              c.chunk_index,
              c.metadata_json,
              d.id AS document_id,
              d.file_name,
              1 - (c.embedding <=> CAST(:query_embedding AS vector)) AS score
            FROM knowledge_chunks c
            JOIN documents d ON d.id = c.document_id
            WHERE c.knowledge_base_id = :kb_id
              AND c.embedding IS NOT NULL
              AND c.status = 'indexed'
              AND d.deleted_at IS NULL
            ORDER BY c.embedding <=> CAST(:query_embedding AS vector), c.id DESC
            LIMIT :limit
            """
        ),
        {
            "kb_id": knowledge_base["id"],
            "query_embedding": _vector_to_pgvector(query_embedding),
            "limit": top_k,
        },
    )
    rows = [dict(row) for row in result.mappings()]
    return [
        _chunk_response(item, retrieval_mode="vector")
        for item in rows
        if float(item.get("score") or 0.0) >= score_threshold
    ]


async def _retrieve_chunks_by_keyword(
    conn,
    *,
    knowledge_base_id: int,
    query: str,
    top_k: int,
    score_threshold: float,
) -> list[dict[str, Any]]:
    result = await conn.execute(
        text(
            """
            SELECT
              c.id,
              c.content,
              c.chunk_index,
              c.metadata_json,
              d.id AS document_id,
              d.file_name
            FROM knowledge_chunks c
            JOIN documents d ON d.id = c.document_id
            WHERE c.knowledge_base_id = :kb_id
              AND c.status != 'deleted'
              AND d.deleted_at IS NULL
            ORDER BY c.created_at DESC, c.id DESC
            LIMIT :limit
            """
        ),
        {"kb_id": knowledge_base_id, "limit": max(top_k * 5, top_k)},
    )
    candidates = [dict(row) for row in result.mappings()]
    return [
        _chunk_response(item, retrieval_mode="keyword")
        for item in rank_chunks(candidates, query)
        if item["score"] >= score_threshold
    ][:top_k]


async def process_document_job(conn, job: dict[str, Any]) -> int:
    document = await _get_document_for_processing(conn, job["document_id"])
    knowledge_base = await _get_knowledge_base_for_processing(conn, document["knowledge_base_id"])
    await _mark_document_status(conn, document["id"], "parsing")
    extracted_text = extract_text_from_file(document["storage_url"], document.get("file_type"))

    await _mark_document_status(conn, document["id"], "chunking")
    chunks = chunk_text(extracted_text)
    if not chunks:
        raise DocumentProcessingError("chunking", "document produced no text chunks")

    await _mark_document_status(conn, document["id"], "embedding")
    try:
        embeddings = await embed_texts(conn, knowledge_base, chunks)
    except DocumentProcessingError:
        raise
    except Exception as exc:
        raise DocumentProcessingError("embedding", str(exc)) from exc

    await _replace_document_chunks(conn, document, chunks, embeddings, knowledge_base)
    await _mark_document_status(conn, document["id"], "indexed")
    return len(chunks)


async def mark_document_failed(
    conn,
    document_id: int,
    *,
    error_stage: str,
    error_message: str,
) -> None:
    await _mark_document_status(
        conn,
        document_id,
        "failed",
        error_stage=error_stage,
        error_message=error_message,
    )


async def _get_document_for_processing(conn, document_id: int) -> dict[str, Any]:
    result = await conn.execute(
        text(
            """
            SELECT *
            FROM documents
            WHERE id = :document_id AND deleted_at IS NULL
            """
        ),
        {"document_id": document_id},
    )
    row = result.mappings().one_or_none()
    if row is None:
        raise DocumentProcessingError("parse", "document not found")
    document = dict(row)
    if not document.get("storage_url"):
        raise DocumentProcessingError("parse", "document has no storage_url")
    return document


async def _get_knowledge_base_for_processing(conn, knowledge_base_id: int) -> dict[str, Any]:
    result = await conn.execute(
        text(
            """
            SELECT *
            FROM knowledge_bases
            WHERE id = :knowledge_base_id AND deleted_at IS NULL AND status != 'deleted'
            """
        ),
        {"knowledge_base_id": knowledge_base_id},
    )
    row = result.mappings().one_or_none()
    if row is None:
        raise DocumentProcessingError("embedding", "knowledge base not found")
    knowledge_base = dict(row)
    knowledge_base["embedding_dim"] = int(
        knowledge_base.get("embedding_dim") or DEFAULT_EMBEDDING_DIM
    )
    return knowledge_base


async def _mark_document_status(
    conn,
    document_id: int,
    status: str,
    *,
    error_stage: str | None = None,
    error_message: str | None = None,
) -> None:
    await conn.execute(
        text(
            """
            UPDATE documents
            SET status = :status,
                error_stage = :error_stage,
                error_message = :error_message,
                updated_at = now()
            WHERE id = :document_id
            """
        ),
        {
            "document_id": document_id,
            "status": status,
            "error_stage": error_stage,
            "error_message": error_message,
        },
    )


async def _replace_document_chunks(
    conn,
    document: dict[str, Any],
    chunks: list[str],
    embeddings: list[list[float]],
    knowledge_base: dict[str, Any],
) -> None:
    if len(chunks) != len(embeddings):
        raise DocumentProcessingError("embedding", "embedding count does not match chunk count")

    await conn.execute(
        text("DELETE FROM knowledge_chunks WHERE document_id = :document_id"),
        {"document_id": document["id"]},
    )
    statement = text(
        """
        INSERT INTO knowledge_chunks (
          knowledge_base_id,
          document_id,
          chunk_index,
          content,
          token_count,
          embedding,
          status,
          metadata_json
        )
        VALUES (
          :knowledge_base_id,
          :document_id,
          :chunk_index,
          :content,
          :token_count,
          CAST(:embedding AS vector),
          'indexed',
          :metadata_json
        )
        """
    ).bindparams(bindparam("metadata_json", type_=JSONB))
    provider, model = _embedding_provider_and_model(knowledge_base)
    for index, chunk in enumerate(chunks):
        await conn.execute(
            statement,
            {
                "knowledge_base_id": document["knowledge_base_id"],
                "document_id": document["id"],
                "chunk_index": index,
                "content": chunk,
                "token_count": _estimate_token_count(chunk),
                "embedding": _vector_to_pgvector(embeddings[index]),
                "metadata_json": {
                    "source_file_name": document["file_name"],
                    "source_file_type": document.get("file_type"),
                    "embedding_provider": provider,
                    "embedding_model": model,
                    "embedding_dim": len(embeddings[index]),
                },
            },
        )


async def embed_texts(
    conn,
    knowledge_base: dict[str, Any],
    texts: list[str],
) -> list[list[float]]:
    if not texts:
        return []

    provider, model = _embedding_provider_and_model(knowledge_base)
    if provider == "openai":
        api_key = await resolve_openai_api_key(conn)
        if not api_key:
            raise DocumentProcessingError(
                "embedding",
                "OpenAI API key is required for embedding_provider=openai",
            )
        return await _embed_texts_with_openai(api_key, model, texts, knowledge_base)

    dimension = int(knowledge_base.get("embedding_dim") or DEFAULT_EMBEDDING_DIM)
    return [local_hash_embedding(text_value, dimension=dimension) for text_value in texts]


def local_hash_embedding(text_value: str, *, dimension: int = DEFAULT_EMBEDDING_DIM) -> list[float]:
    if dimension <= 0:
        raise ValueError("embedding dimension must be greater than 0")

    vector = [0.0] * dimension
    tokens = _embedding_tokens(text_value)
    if not tokens:
        tokens = [""]

    for token in tokens:
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % dimension
        sign = 1.0 if digest[4] % 2 == 0 else -1.0
        vector[index] += sign

    norm = math.sqrt(sum(value * value for value in vector))
    if norm == 0:
        return vector
    return [round(value / norm, 8) for value in vector]


def _embedding_provider_and_model(knowledge_base: dict[str, Any]) -> tuple[str, str]:
    config = knowledge_base.get("config_json") or {}
    model = str(
        knowledge_base.get("embedding_model")
        or config.get("embedding_model")
        or "local-hash"
    )
    provider = str(config.get("embedding_provider") or config.get("provider") or "").lower()
    if not provider:
        provider = "local" if model in LOCAL_EMBEDDING_MODELS else "openai"
    if provider == "local-hash":
        provider = "local"
    return provider, model


async def _embed_texts_with_openai(
    api_key: str,
    model: str,
    texts: list[str],
    knowledge_base: dict[str, Any],
) -> list[list[float]]:
    try:
        from openai import AsyncOpenAI
    except ImportError as exc:
        raise DocumentProcessingError("embedding", "openai package is not installed") from exc

    client = AsyncOpenAI(api_key=api_key)
    response = await client.embeddings.create(model=model, input=texts)
    embeddings = [list(item.embedding) for item in response.data]
    expected_dim = int(knowledge_base.get("embedding_dim") or DEFAULT_EMBEDDING_DIM)
    for embedding in embeddings:
        if len(embedding) != expected_dim:
            raise DocumentProcessingError(
                "embedding",
                f"embedding dimension mismatch: expected {expected_dim}, got {len(embedding)}",
            )
    return embeddings


def _embedding_tokens(text_value: str) -> list[str]:
    normalized = _normalize_text(text_value).lower()
    return [token for token in normalized.replace("\n", " ").split(" ") if token]


def _vector_to_pgvector(vector: list[float]) -> str:
    return "[" + ",".join(f"{value:.8f}" for value in vector) + "]"


def _chunk_response(item: dict[str, Any], *, retrieval_mode: str) -> dict[str, Any]:
    return {
        "chunk_id": str(item["id"]),
        "content": item["content"],
        "score": float(item.get("score") or 0.0),
        "retrieval_mode": retrieval_mode,
        "source": {
            "document_id": item["document_id"],
            "file_name": item["file_name"],
            "chunk_index": item["chunk_index"],
            "metadata_json": item.get("metadata_json") or {},
        },
    }


def _extract_json_text(file_path: Path) -> str:
    with file_path.open("r", encoding="utf-8") as file:
        value = json.load(file)
    return "\n".join(_walk_json_text(value))


def _extract_csv_text(file_path: Path) -> str:
    content = _read_utf8_text(file_path, "CSV")
    rows = csv.reader(StringIO(content))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _extract_plain_text(file_path: Path) -> str:
    return _read_utf8_text(file_path, "plain text")


def _extract_markdown_text(file_path: Path) -> str:
    content = _read_utf8_text(file_path, "Markdown")
    try:
        from markdown_it import MarkdownIt
    except ImportError:
        return _fallback_markdown_text(content)

    tokens = MarkdownIt("commonmark").parse(content)
    pieces: list[str] = []
    for token in tokens:
        if token.type in {"text", "code_inline", "code_block", "fence"}:
            pieces.append(token.content)
        elif token.type in {"softbreak", "hardbreak"}:
            pieces.append("\n")
        elif token.type == "inline" and token.children:
            for child in token.children:
                if child.type in {"text", "code_inline"}:
                    pieces.append(child.content)
                elif child.type in {"softbreak", "hardbreak"}:
                    pieces.append("\n")
        elif token.type in {"paragraph_close", "heading_close", "list_item_close"}:
            pieces.append("\n")

    return _normalize_text("".join(pieces))


def _extract_pdf_text(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentProcessingError("parse", "pypdf is not installed") from exc

    try:
        reader = PdfReader(str(file_path))
        return _normalize_text("\n".join(page.extract_text() or "" for page in reader.pages))
    except Exception as exc:
        raise DocumentProcessingError("parse", f"failed to extract PDF text: {exc}") from exc


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentProcessingError("parse", "python-docx is not installed") from exc

    try:
        document = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return _normalize_text("\n".join([*paragraphs, *table_rows]))
    except Exception as exc:
        raise DocumentProcessingError("parse", f"failed to extract DOCX text: {exc}") from exc


def _walk_json_text(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, int | float | bool):
        return [str(value)]
    if isinstance(value, list):
        items: list[str] = []
        for item in value:
            items.extend(_walk_json_text(item))
        return items
    if isinstance(value, dict):
        items = []
        for key, item in value.items():
            child_text = " ".join(_walk_json_text(item)).strip()
            items.append(f"{key}: {child_text}" if child_text else str(key))
        return items
    return [str(value)]


def _normalize_text(text_value: str) -> str:
    return "\n".join(line.rstrip() for line in text_value.replace("\r\n", "\n").split("\n")).strip()


def _estimate_token_count(text_value: str) -> int:
    return max(1, len(text_value.split()))


def _is_text_content_type(file_type: str | None) -> bool:
    return bool(
        file_type
        and file_type.startswith("text/")
        and file_type not in {"text/csv", "text/markdown"}
    )


def _is_docx_content_type(file_type: str | None) -> bool:
    return file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _is_markdown_content_type(file_type: str | None) -> bool:
    return file_type in {"text/markdown", "text/x-markdown"}


def _read_utf8_text(file_path: Path, label: str) -> str:
    try:
        return file_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentProcessingError(
            "parse",
            f"failed to decode {label} as UTF-8 text: {exc.reason}",
        ) from exc


def _fallback_markdown_text(content: str) -> str:
    lines: list[str] = []
    in_fence = False
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
            continue
        if in_fence:
            lines.append(line)
            continue
        line = stripped.lstrip("#").strip()
        line = line.lstrip(">-*+0123456789. ").strip()
        line = line.replace("**", "").replace("__", "").replace("`", "")
        if line:
            lines.append(line)
    return _normalize_text("\n".join(lines))
